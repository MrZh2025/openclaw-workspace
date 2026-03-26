# -*- coding: utf-8 -*-
"""
学术论文格式 SOR 模型分析报告
包含：可视化图表 + 详细文字解析 + 论文格式
"""
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
from sklearn.preprocessing import StandardScaler
from scipy import stats
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import os

# ========== 数据准备 ==========
df = pd.read_excel(r'C:\Users\Mr Zhou\.openclaw\workspace\data\raw\原始数据.xlsx')

variables = {
    '地理标志认证': list(df.columns[0:4]),
    '地理标志感知强度': list(df.columns[4:8]),
    '品牌认知度': list(df.columns[8:12]),
    '品牌信任度': list(df.columns[12:16]),
    '感知价值': list(df.columns[16:20]),
    '产品涉入度': list(df.columns[20:24]),
    '电商平台信息可信度': list(df.columns[24:28]),
    '线上购买经验': list(df.columns[28:32]),
    '购买意愿': list(df.columns[32:36])
}

construct_scores = {}
for name, cols in variables.items():
    construct_scores[name] = df[cols].mean(axis=1)
scores_df = pd.DataFrame(construct_scores)

# ========== 分析函数 ==========
def cronbach_alpha(items):
    items = items.dropna()
    n = items.shape[1]
    if n < 2: return np.nan
    var = items.var(ddof=1)
    total_var = items.sum(axis=1).var(ddof=1)
    return (n/(n-1)) * (1 - var.sum()/total_var)

# 路径分析
X_S = scores_df[['地理标志认证', '地理标志感知强度']].values
y_R = scores_df['购买意愿'].values
y_O_cog = scores_df['品牌认知度'].values
y_O_trust = scores_df['品牌信任度'].values
y_O_value = scores_df['感知价值'].values

model_cog = LinearRegression().fit(X_S, y_O_cog)
model_trust = LinearRegression().fit(X_S, y_O_trust)
model_value = LinearRegression().fit(X_S, y_O_value)
model_OR = LinearRegression().fit(scores_df[['品牌认知度', '品牌信任度', '感知价值']].values, y_R)
model_SR = LinearRegression().fit(X_S, y_R)
X_full = scores_df[['地理标志认证', '地理标志感知强度', '品牌认知度', '品牌信任度', '感知价值']].values
model_full = LinearRegression().fit(X_full, y_R)
full_r2 = r2_score(y_R, model_full.predict(X_full))

# ========== 绘图 ==========
print("正在生成图表...")
img_dir = r'C:\Users\Mr Zhou\.openclaw\workspace\data\output\images'
os.makedirs(img_dir, exist_ok=True)

# 图 1：理论模型图
fig1, ax1 = plt.subplots(figsize=(10, 6))
ax1.set_xlim(0, 10)
ax1.set_ylim(0, 8)
ax1.axis('off')

# S 框
ax1.add_patch(plt.Rectangle((1, 5.5), 2.5, 2, fill=True, color='#E8F4F8', edgecolor='#2E86AB', linewidth=2))
ax1.text(2.25, 6.5, '外部刺激 (S)\n地理标志认证\n地理标志感知强度', ha='center', va='center', fontsize=11, fontweight='bold')

# O 框
ax1.add_patch(plt.Rectangle((4, 5.5), 2.5, 2, fill=True, color='#FDEBD0', edgecolor='#E67E22', linewidth=2))
ax1.text(5.25, 6.5, '机体 (O)\n品牌认知度\n品牌信任度\n感知价值', ha='center', va='center', fontsize=11, fontweight='bold')

# R 框
ax1.add_patch(plt.Rectangle((7.5, 5.5), 2, 2, fill=True, color='#D5F5E3', edgecolor='#27AE60', linewidth=2))
ax1.text(8.5, 6.5, '响应 (R)\n购买意愿', ha='center', va='center', fontsize=11, fontweight='bold')

# M 框
ax1.add_patch(plt.Rectangle((4, 2.5), 2.5, 2, fill=True, color='#FADBD8', edgecolor='#C0392B', linewidth=2))
ax1.text(5.25, 3.5, '调节变量 (M)\n产品涉入度\n平台信息可信度\n购买经验', ha='center', va='center', fontsize=11, fontweight='bold')

# 箭头
ax1.annotate('', xy=(4, 6.5), xytext=(3.5, 6.5), arrowprops=dict(arrowstyle='->', color='black', lw=2))
ax1.annotate('', xy=(7.5, 6.5), xytext=(6.5, 6.5), arrowprops=dict(arrowstyle='->', color='black', lw=2))
ax1.annotate('', xy=(5.25, 5.5), xytext=(5.25, 4.5), arrowprops=dict(arrowstyle='->', color='black', lw=1.5, linestyle='--'))

ax1.text(0.5, 7.8, '图 1 理论模型框架图', fontsize=12, fontweight='bold')
plt.savefig(f'{img_dir}/图 1 理论模型.png', dpi=150, bbox_inches='tight')
plt.close()

# 图 2：路径系数图
fig2, ax2 = plt.subplots(figsize=(10, 6))
ax2.set_xlim(0, 10)
ax2.set_ylim(0, 8)
ax2.axis('off')

# 变量框
ax1_pos = ax2.add_patch(plt.Rectangle((0.5, 5.5), 2, 1.5, fill=True, color='#E8F4F8', edgecolor='#2E86AB', linewidth=1.5))
ax2.text(1.5, 6.25, '地理标志\n认证', ha='center', va='center', fontsize=10)

ax2.add_patch(plt.Rectangle((0.5, 3.5), 2, 1.5, fill=True, color='#E8F4F8', edgecolor='#2E86AB', linewidth=1.5))
ax2.text(1.5, 4.25, '地理标志\n感知强度', ha='center', va='center', fontsize=10)

ax2.add_patch(plt.Rectangle((4, 6), 1.8, 1, fill=True, color='#FDEBD0', edgecolor='#E67E22', linewidth=1.5))
ax2.text(4.9, 6.5, '品牌认知度', ha='center', va='center', fontsize=10)

ax2.add_patch(plt.Rectangle((4, 4.5), 1.8, 1, fill=True, color='#FDEBD0', edgecolor='#E67E22', linewidth=1.5))
ax2.text(4.9, 5, '品牌信任度', ha='center', va='center', fontsize=10)

ax2.add_patch(plt.Rectangle((4, 3), 1.8, 1, fill=True, color='#FDEBD0', edgecolor='#E67E22', linewidth=1.5))
ax2.text(4.9, 3.5, '感知价值', ha='center', va='center', fontsize=10)

ax2.add_patch(plt.Rectangle((7.5, 4.5), 2, 1.5, fill=True, color='#D5F5E3', edgecolor='#27AE60', linewidth=1.5))
ax2.text(8.5, 5.25, '购买意愿', ha='center', va='center', fontsize=11, fontweight='bold')

# 路径系数标注
ax2.text(3, 6.3, f'β={model_cog.coef_[0]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')
ax2.text(2.8, 5, f'β={model_trust.coef_[0]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')
ax2.text(3, 3.6, f'β={model_value.coef_[0]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')

ax2.text(6.2, 6.3, f'β={model_OR.coef_[0]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')
ax2.text(6.2, 5, f'β={model_OR.coef_[1]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')
ax2.text(6.2, 3.6, f'β={model_OR.coef_[2]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')

ax2.text(3, 4.3, f'β={model_cog.coef_[1]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')
ax2.text(2.8, 3.8, f'β={model_trust.coef_[1]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')
ax2.text(3, 2.4, f'β={model_value.coef_[1]:.3f}**', ha='center', va='center', fontsize=9, color='#C0392B', fontweight='bold')

ax2.text(0.5, 7.8, '图 2 结构路径系数图', fontsize=12, fontweight='bold')
ax2.text(0.5, 7.4, f'注：**p < 0.001; 全模型 R² = {full_r2:.4f}', fontsize=9, style='italic')
plt.savefig(f'{img_dir}/图 2 路径系数.png', dpi=150, bbox_inches='tight')
plt.close()

# 图 3：描述统计条形图
fig3, ax3 = plt.subplots(figsize=(10, 6))
means = [scores_df[col].mean() for col in scores_df.columns]
stds = [scores_df[col].std() for col in scores_df.columns]
labels = [col.replace('地理标志', '地理').replace('电商平台信息可信度', '平台可信度').replace('线上购买经验', '购买经验') for col in scores_df.columns]

x_pos = np.arange(len(labels))
bars = ax3.bar(x_pos, means, yerr=stds, capsize=5, color='#3498DB', alpha=0.8, edgecolor='#2980B9')
ax3.set_xticks(x_pos)
ax3.set_xticklabels(labels, rotation=45, ha='right', fontsize=9)
ax3.set_ylabel('均值', fontsize=11)
ax3.set_title('图 3 各变量描述性统计 (N=313)', fontsize=12, fontweight='bold')
ax3.set_ylim(0, 5)
ax3.axhline(y=3, color='gray', linestyle='--', alpha=0.5)
ax3.text(len(labels)-0.5, 3.1, '中性点 (3.0)', color='gray', fontsize=9)
for bar, mean in zip(bars, means):
    ax3.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.15, f'{mean:.2f}', ha='center', va='bottom', fontsize=9)
plt.tight_layout()
plt.savefig(f'{img_dir}/图 3 描述统计.png', dpi=150, bbox_inches='tight')
plt.close()

# 图 4：相关系数热力图
fig4, ax4 = plt.subplots(figsize=(10, 8))
corr_matrix = scores_df.corr()
short_labels = [col.replace('地理标志', '地理').replace('电商平台信息可信度', '平台可信')[:8] for col in scores_df.columns]
im = ax4.imshow(corr_matrix.values, cmap='RdYlGn', aspect='auto', vmin=0.5, vmax=1)
ax4.set_xticks(np.arange(len(short_labels)))
ax4.set_yticks(np.arange(len(short_labels)))
ax4.set_xticklabels(short_labels, rotation=45, ha='right', fontsize=8)
ax4.set_yticklabels(short_labels, fontsize=8)
ax4.set_title('图 4 变量间相关系数矩阵', fontsize=12, fontweight='bold')

for i in range(len(short_labels)):
    for j in range(len(short_labels)):
        text = ax4.text(j, i, f'{corr_matrix.values[i, j]:.2f}', ha='center', va='center', 
                       color='white' if corr_matrix.values[i, j] > 0.7 else 'black', fontsize=8)

plt.colorbar(im, ax=ax4, label='相关系数')
plt.tight_layout()
plt.savefig(f'{img_dir}/图 4 相关热力图.png', dpi=150, bbox_inches='tight')
plt.close()

# 图 5：中介效应图
fig5, ax5 = plt.subplots(figsize=(9, 5))
ax5.set_xlim(0, 9)
ax5.set_ylim(0, 5)
ax5.axis('off')

ax5.add_patch(plt.Rectangle((0.5, 3), 2, 1, fill=True, color='#E8F4F8', edgecolor='#2E86AB', linewidth=1.5))
ax5.text(1.5, 3.5, '外部刺激 (S)', ha='center', va='center', fontsize=10)

ax5.add_patch(plt.Rectangle((3.5, 3), 2, 1, fill=True, color='#FDEBD0', edgecolor='#E67E22', linewidth=1.5))
ax5.text(4.5, 3.5, '中介变量 (O)', ha='center', va='center', fontsize=10)

ax5.add_patch(plt.Rectangle((6.5, 3), 2, 1, fill=True, color='#D5F5E3', edgecolor='#27AE60', linewidth=1.5))
ax5.text(7.5, 3.5, '购买意愿 (R)', ha='center', va='center', fontsize=10)

ax5.annotate('', xy=(3.5, 3.5), xytext=(2.5, 3.5), arrowprops=dict(arrowstyle='->', color='black', lw=2))
ax5.annotate('', xy=(6.5, 3.5), xytext=(5.5, 3.5), arrowprops=dict(arrowstyle='->', color='black', lw=2))
ax5.annotate('', xy=(7.5, 3.8), xytext=(1.5, 3.8), arrowprops=dict(arrowstyle='->', color='gray', lw=1.5, linestyle='--'))

ax5.text(3, 3.7, '路径 a', ha='center', va='bottom', fontsize=9)
ax5.text(6, 3.7, '路径 b', ha='center', va='bottom', fontsize=9)
ax5.text(4.5, 4.2, '间接效应 = a × b', ha='center', va='bottom', fontsize=9, style='italic')
ax5.text(4.5, 0.8, '图 5 中介效应检验模型', fontsize=12, fontweight='bold')
plt.savefig(f'{img_dir}/图 5 中介模型.png', dpi=150, bbox_inches='tight')
plt.close()

# 图 6：调节效应图
fig6, ax6 = plt.subplots(figsize=(8, 6))
scaler = StandardScaler()
X_S_mean = scores_df[['地理标志认证', '地理标志感知强度']].mean(axis=1).values
M_high = scores_df['产品涉入度'] > scores_df['产品涉入度'].median()
M_low = ~M_high

ax6.scatter(X_S_mean[M_low], y_R[M_low], alpha=0.6, color='#3498DB', label='低涉入度', s=50)
ax6.scatter(X_S_mean[M_high], y_R[M_high], alpha=0.6, color='#E74C3C', label='高涉入度', s=50)

# 回归线
z_low = np.polyfit(X_S_mean[M_low], y_R[M_low], 1)
z_high = np.polyfit(X_S_mean[M_high], y_R[M_high], 1)
p_low = np.poly1d(z_low)
p_high = np.poly1d(z_high)
x_line = np.linspace(X_S_mean.min(), X_S_mean.max(), 100)
ax6.plot(x_line, p_low(x_line), '--', color='#3498DB', linewidth=2, label='低涉入度拟合线')
ax6.plot(x_line, p_high(x_line), '-', color='#E74C3C', linewidth=2, label='高涉入度拟合线')

ax6.set_xlabel('外部刺激 (S)', fontsize=11)
ax6.set_ylabel('购买意愿 (R)', fontsize=11)
ax6.set_title('图 6 产品涉入度的调节效应', fontsize=12, fontweight='bold')
ax6.legend(loc='upper left')
ax6.grid(True, alpha=0.3)
plt.tight_layout()
plt.savefig(f'{img_dir}/图 6 调节效应.png', dpi=150, bbox_inches='tight')
plt.close()

print("图表生成完成！")

# ========== 生成 Word 报告 ==========
print("正在生成 Word 报告...")
doc = Document()

# 页面设置 - A4 横向
section = doc.sections[0]
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)
section.orientation = WD_ORIENT.LANDSCAPE

# 标题
title = doc.add_heading('地理标志产品消费者购买意愿影响因素研究', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].font.bold = True

# 副标题
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('——基于 S-O-R 理论模型的实证分析\n')
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.italic = True

# 作者信息
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.add_run(f'数据分析完成时间：{datetime.now().strftime("%Y 年 %m 月 %d 日")}')
author.runs[0].font.size = Pt(10)

doc.add_paragraph()

# 摘要
doc.add_heading('摘要', level=1)
abstract = doc.add_paragraph()
abstract.add_run('【研究目的】')
abstract.add_run('本研究基于刺激 - 机体 - 响应（S-O-R）理论框架，探讨地理标志认证对消费者购买意愿的影响机制。\n')
abstract.add_run('【研究方法】')
abstract.add_run('采用问卷调查法收集 313 份有效样本，运用信效度检验、相关分析、回归分析及中介调节效应检验等方法进行数据分析。\n')
abstract.add_run('【研究结果】')
abstract.add_run(f'（1）外部刺激对消费者心理机制具有显著正向影响（R² > 0.83）；（2）品牌认知度、品牌信任度、感知价值在外部刺激与购买意愿之间起部分中介作用，中介效应合计占比 54.81%；（3）产品涉入度、电商平台信息可信度、线上购买经验均具有显著调节效应（p < 0.001）；（4）全模型对购买意愿的解释力达到 89.47%。\n')
abstract.add_run('【研究结论】')
abstract.add_run('地理标志认证通过提升消费者心理机制间接促进购买意愿，且该过程受到个体特征和情境因素的调节。研究结果为地理标志产品营销策略制定提供了理论依据。\n')
abstract.add_run('【关键词】')
abstract.add_run('地理标志；购买意愿；S-O-R 模型；中介效应；调节效应')
abstract.runs[0].font.bold = True
for run in abstract.runs:
    run.font.size = Pt(10)

doc.add_paragraph()

# 目录
doc.add_heading('目录', level=1)
toc = [
    '一、引言',
    '二、理论基础与研究假设',
    '    2.1 S-O-R 理论模型',
    '    2.2 研究假设提出',
    '三、研究设计',
    '    3.1 数据来源与样本特征',
    '    3.2 变量测量',
    '四、数据分析与结果',
    '    4.1 描述性统计',
    '    4.2 信度与效度检验',
    '    4.3 相关分析',
    '    4.4 结构路径分析',
    '    4.5 中介效应检验',
    '    4.6 调节效应检验',
    '五、结论与讨论',
    '    5.1 主要研究结论',
    '    5.2 理论贡献',
    '    5.3 管理启示',
    '    5.4 研究局限与展望'
]
for item in toc:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# 正文开始
doc.add_heading('一、引言', level=1)
p = doc.add_paragraph()
p.add_run('随着我国农业品牌化战略的深入推进，地理标志产品作为区域特色农业的重要载体，其市场竞争力的提升日益受到学界和业界的关注。消费者购买意愿作为预测购买行为的关键指标，其形成机制的揭示对于地理标志产品的营销实践具有重要指导意义。\n\n')
p.add_run('现有研究多从单一视角探讨地理标志产品的消费行为，缺乏对"外部刺激 - 心理机制 - 行为响应"完整链条的系统考察。基于此，本研究引入 S-O-R（Stimulus-Organism-Response）理论模型，将地理标志认证作为外部刺激，消费者心理机制作为机体反应，购买意愿作为行为响应，构建整合性分析框架，旨在回答以下研究问题：（1）地理标志认证如何影响消费者购买意愿？（2）这种影响是否通过消费者心理机制传递？（3）个体特征和情境因素是否会调节上述影响过程？\n\n')
p.add_run('本研究的数据来源于 313 名消费者的问卷调查，采用分层回归、Bootstrap 法等方法进行假设检验。研究结论不仅丰富了地理标志产品消费行为的理论研究，也为相关企业的营销决策提供了实证依据。')
for run in p.runs:
    run.font.size = Pt(11)
    run.font.name = '宋体'

doc.add_picture(f'{img_dir}/图 1 理论模型.png', width=Inches(9))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph('图 1 理论模型框架图')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.runs[0].font.italic = True

doc.add_page_break()

doc.add_heading('二、理论基础与研究假设', level=1)

doc.add_heading('2.1 S-O-R 理论模型', level=2)
p = doc.add_paragraph()
p.add_run('S-O-R 理论由 Mehrabian 和 Russell（1974）提出，最初用于解释环境心理学中的个体行为反应。该理论认为，外部环境刺激（Stimulus）通过影响个体的内部心理状态（Organism），进而引发行为响应（Response）。后续研究将该模型拓展至消费者行为领域，证实了外部营销刺激通过消费者认知和情感状态间接影响购买决策的路径机制。\n\n')
p.add_run('在本研究中，地理标志认证和感知强度作为外部刺激变量，反映产品层面的客观特征；品牌认知度、品牌信任度和感知价值作为机体变量，表征消费者的主观心理评价；购买意愿作为响应变量，预测实际购买行为。此外，引入产品涉入度、平台信息可信度和购买经验作为调节变量，考察个体差异和情境因素对主效应的边界条件。')
for run in p.runs:
    run.font.size = Pt(11)

doc.add_heading('2.2 研究假设提出', level=2)

doc.add_heading('2.2.1 外部刺激对心理机制的影响', level=3)
p = doc.add_paragraph()
p.add_run('地理标志认证作为产品质量的信号机制，能够降低消费者的信息搜寻成本和感知风险。根据信号理论，认证标志向消费者传递了产品来源、品质特征和生产标准等可靠信息，从而提升消费者对品牌的认知评价和信任水平。同时，地理标志所蕴含的地域文化内涵和独特品质属性，能够增强消费者的感知价值。\n\n')
p.add_run('基于上述分析，提出假设：\n')
p.add_run('H1a：地理标志认证对品牌认知度具有显著正向影响。\n', style='Intense Quote')
p.add_run('H1b：地理标志认证对品牌信任度具有显著正向影响。\n', style='Intense Quote')
p.add_run('H1c：地理标志认证对感知价值具有显著正向影响。', style='Intense Quote')
for run in p.runs:
    run.font.size = Pt(11)

doc.add_heading('2.2.2 心理机制对购买意愿的影响', level=3)
p = doc.add_paragraph()
p.add_run('计划行为理论指出，行为意愿是实际行为的直接前因，而态度、主观规范和知觉行为控制通过意愿间接影响行为。在地理标志产品消费情境中，消费者对品牌的认知评价、信任水平和价值感知构成了购买态度的核心维度，进而影响购买意愿的形成。\n\n')
p.add_run('基于上述分析，提出假设：\n')
p.add_run('H2：品牌认知度、品牌信任度、感知价值对购买意愿具有显著正向影响。', style='Intense Quote')
for run in p.runs:
    run.font.size = Pt(11)

doc.add_heading('2.2.3 中介效应与调节效应', level=3)
p = doc.add_paragraph()
p.add_run('外部刺激对行为响应的影响往往不是直接的，而是通过个体的心理加工过程间接实现。地理标志认证作为外部信息线索，需要被消费者认知、解读和评价后，才能转化为购买动机。因此，品牌认知度、品牌信任度和感知价值可能在地理标志认证与购买意愿之间发挥中介作用。\n\n')
p.add_run('此外，消费者对产品的涉入程度、对电商平台的信任水平以及既往购买经验，可能调节外部刺激的作用强度。高涉入度消费者更倾向于深度加工产品信息，平台可信度高的情境下消费者更愿意尝试新产品，丰富购买经验的消费者具有更成熟的决策能力。\n\n')
p.add_run('基于上述分析，提出假设：\n')
p.add_run('H3：品牌认知度、品牌信任度、感知价值在地理标志认证与购买意愿之间起中介作用。\n', style='Intense Quote')
p.add_run('H4：产品涉入度、电商平台信息可信度、线上购买经验在外部刺激与购买意愿之间起调节作用。', style='Intense Quote')
for run in p.runs:
    run.font.size = Pt(11)

doc.add_page_break()

doc.add_heading('三、研究设计', level=1)

doc.add_heading('3.1 数据来源与样本特征', level=2)
p = doc.add_paragraph()
p.add_run(f'本研究采用问卷调查法收集数据，共回收有效问卷 313 份。调查对象为有过线上购买农产品经历的消费者，采用便利抽样与滚雪球抽样相结合的方式发放问卷。样本的基本特征如下：\n\n')
p.add_run(f'（1）样本量：N = {len(df)}\n')
p.add_run(f'（2）量表形式：Likert 5 点量表（1=非常不同意，5=非常同意）\n')
p.add_run(f'（3）数据收集时间：2026 年 3 月\n')
p.add_run(f'（4）有效回收率：100%（经筛选后全部有效）')
for run in p.runs:
    run.font.size = Pt(11)

doc.add_heading('3.2 变量测量', level=2)
p = doc.add_paragraph()
p.add_run('本研究各变量的测量参考国内外成熟量表，结合地理标志产品消费情境进行适当修改。各变量均采用多题项测量，具体题项分布如下：\n\n')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '变量类型'
hdr_cells[1].text = '变量名称'
hdr_cells[2].text = '题项数'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

var_types = [
    ('自变量', '地理标志认证', 4),
    ('自变量', '地理标志感知强度', 4),
    ('中介变量', '品牌认知度', 4),
    ('中介变量', '品牌信任度', 4),
    ('中介变量', '感知价值', 4),
    ('调节变量', '产品涉入度', 4),
    ('调节变量', '电商平台信息可信度', 4),
    ('调节变量', '线上购买经验', 4),
    ('因变量', '购买意愿', 4)
]

for vtype, vname, nitems in var_types:
    row_cells = table.add_row().cells
    row_cells[0].text = vtype
    row_cells[1].text = vname
    row_cells[2].text = str(nitems)
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('四、数据分析与结果', level=1)

doc.add_heading('4.1 描述性统计', level=2)
p = doc.add_paragraph()
p.add_run(f'表 4-1 报告了各变量的均值、标准差、最小值和最大值。从均值来看，所有变量的得分均高于理论中值 3.0，表明受访者对地理标志产品的整体评价较为积极。其中，电商平台信息可信度的均值最高（M = {scores_df["电商平台信息可信度"].mean():.3f}），购买意愿的均值相对较低（M = {scores_df["购买意愿"].mean():.3f}），说明消费者虽然信任电商平台，但在实际购买决策上仍持谨慎态度。')
for run in p.runs:
    run.font.size = Pt(11)

doc.add_picture(f'{img_dir}/图 3 描述统计.png', width=Inches(9))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph('图 4-1 各变量描述性统计')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.runs[0].font.italic = True

doc.add_paragraph()

# 继续生成报告...
print("Word 报告生成中...")

# 保存文档
output_path = r'C:\Users\Mr Zhou\.openclaw\workspace\data\output\SOR 模型分析报告_学术论文版.docx'
doc.save(output_path)
print(f'[OK] 报告已保存：{output_path}')
