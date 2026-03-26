from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_radiation_license_report():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('辐射安全许可证办理分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('办理难点、周期与卡点全面解析')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 空行
    
    # ========== 第一部分：整体办理周期 ==========
    doc.add_heading('一、整体办理周期概览', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    headers = ['环节', '理论时限', '实际耗时', '说明']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # 数据
    data = [
        ['前期准备', '-', '1-3 个月', '材料收集、人员培训'],
        ['环评报告', '30 工作日', '2-6 个月', '最大瓶颈环节'],
        ['设计审查', '20 工作日', '1-2 个月', '防护设施设计'],
        ['技术评审', '30 工作日', '1-3 个月', '专家评审会'],
        ['现场核查', '10 工作日', '2-4 周', '设施人员到位'],
        ['审批发证', '20 工作日', '1-2 个月', '最终审批'],
        ['合计', '约 4 个月', '6-18 个月', '视项目复杂程度'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # ========== 第二部分：最耽误时间的环节 ==========
    doc.add_heading('二、最耽误时间的 3 个环节', level=1)
    
    doc.add_heading('1️⃣ 环境影响评价（环评）— 最大瓶颈', level=2)
    p = doc.add_paragraph()
    p.add_run('预计耗时：').bold = True
    p.add_run('2-6 个月\n')
    p.add_run('主要难点：\n').bold = True
    p.add_run('• 需要委托有资质的环评机构编制报告\n')
    p.add_run('• 涉及辐射剂量评估、周边环境敏感点分析\n')
    p.add_run('• 公众参与环节（公示 10-15 工作日，易被投诉）\n')
    p.add_run('• 环保部门审批排队时间长\n')
    p.add_run('• 如需编制报告书（非报告表），时间更长\n')
    
    p = doc.add_paragraph()
    p.add_run('常见卡点：\n').bold = True
    p.add_run('❌ 环评机构档期紧张（常见 1-2 个月排队）\n')
    p.add_run('❌ 周边敏感点（学校、医院、居民区）需额外论证\n')
    p.add_run('❌ 公示期间收到反对意见需重新评估\n')
    
    doc.add_heading('2️⃣ 辐射防护设施设计与验收', level=2)
    p = doc.add_paragraph()
    p.add_run('预计耗时：').bold = True
    p.add_run('1-3 个月\n')
    p.add_run('主要难点：\n').bold = True
    p.add_run('• 防护设计需符合 GB 系列标准（如 GB 18871）\n')
    p.add_run('• 屏蔽计算需专业资质单位出具\n')
    p.add_run('• 防护材料（铅板、混凝土等）需检测报告\n')
    p.add_run('• 竣工后需第三方检测验收\n')
    
    doc.add_heading('3️⃣ 技术评审与现场核查', level=2)
    p = doc.add_paragraph()
    p.add_run('预计耗时：').bold = True
    p.add_run('1-3 个月\n')
    p.add_run('主要难点：\n').bold = True
    p.add_run('• 专家评审会需协调多位专家时间\n')
    p.add_run('• 现场核查需所有设施到位、人员持证\n')
    p.add_run('• 整改意见需逐项落实并复验\n')
    
    doc.add_paragraph()
    
    # ========== 第三部分：环评办理详解 ==========
    doc.add_heading('三、环评办理详细分析', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ['阶段', '理论时限', '实际耗时', '说明']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    data = [
        ['委托环评机构', '-', '1-2 周', '比价、合同签订'],
        ['现场踏勘', '5 工作日', '1-2 周', '机构排期 + 踏勘'],
        ['报告编制', '30 工作日', '1-3 个月', '最耗时环节'],
        ['内部审核', '10 工作日', '2-4 周', '环评机构质控'],
        ['报送受理', '5 工作日', '1-2 周', '形式审查'],
        ['技术评估', '20 工作日', '1-2 个月', '专家评审'],
        ['公众参与', '10 工作日', '2-4 周', '公示 + 意见反馈'],
        ['审批决定', '15 工作日', '2-4 周', '生态环境局审批'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # ========== 第四部分：许可类型 ==========
    doc.add_heading('四、辐射安全许可证类型', level=1)
    
    doc.add_heading('按活动类型分类', level=2)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ['类型', '适用范围', '审批层级', '办理周期', '难度']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    data = [
        ['生产放射性同位素', '生产制备同位素', '生态环境部', '12-24 个月', '★★★★★'],
        ['销售放射性同位素', '经营销售同位素', '省级生态环境厅', '6-12 个月', '★★★★'],
        ['使用放射性同位素', '科研医疗工业使用', '省级/市级', '4-10 个月', '★★★'],
        ['生产射线装置', '生产制造射线装置', '省级生态环境厅', '6-12 个月', '★★★★'],
        ['销售射线装置', '经营销售射线装置', '省级/市级', '3-8 个月', '★★★'],
        ['使用射线装置', '诊疗检测科研', '市级/县级', '3-8 个月', '★★'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('按射线装置类别分类', level=2)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ['类别', '风险等级', '典型设备', '环评要求', '办理周期']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    data = [
        ['Ⅰ类', '高', '医用加速器、质子治疗装置', '报告书', '8-18 个月'],
        ['Ⅱ类', '中', 'CT、DR、CR、乳腺机', '报告书', '4-10 个月'],
        ['Ⅲ类', '低', '牙科 X 光机、安检机', '报告表', '2-4 个月'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # ========== 第五部分：加速建议 ==========
    doc.add_heading('五、加速办理建议', level=1)
    
    strategies = [
        ('提前启动环评', '这是最大瓶颈，尽早委托有资质的环评机构，可节省 2-4 周'),
        ('同步准备人员资质', '在环评期间安排人员参加辐射安全培训，可节省 2-3 周'),
        ('选择成熟方案', '参考同类单位的设计，避免反复修改'),
        ('主动沟通审批部门', '提前了解当地审批要求和关注重点，可节省 1-2 周'),
        ('规避敏感点', '选址时避开学校、医院、密集居民区，可节省 1-3 个月'),
        ('选择报告表而非报告书', 'Ⅲ类射线装置可编制报告表，可节省 2-4 个月'),
        ('委托全程服务', '选择可代办审批的环评机构，可节省 1-2 个月'),
    ]
    
    for title, content in strategies:
        p = doc.add_paragraph()
        p.add_run('✅ ').bold = True
        p.add_run(title).bold = True
        p.add_run(f' — {content}')
    
    doc.add_paragraph()
    
    # ========== 第六部分：高风险预警 ==========
    doc.add_heading('六、高风险预警', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ['情况', '风险等级', '建议']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    data = [
        ['周边 50m 内有学校/幼儿园', '🔴 高危', '考虑重新选址'],
        ['位于居民楼内', '🔴 高危', '需额外论证，审批难度大'],
        ['涉及放射源（非射线装置）', '🟡 中危', '需增加安保、监控等措施'],
        ['首次办理无经验', '🟡 中危', '委托专业机构全程代办'],
        ['多设备/多场所', '🟡 中危', '可分批办理，降低单次复杂度'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # ========== 页脚 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    footer.text = '本报告由 OpenClaw 生成 | 仅供参考，具体办理请以当地生态环境部门要求为准'
    
    # 保存文档
    doc.save('C:/Users/Mr Zhou/Desktop/辐射安全许可证办理分析报告.docx')
    print("文档已保存至桌面：辐射安全许可证办理分析报告.docx")

if __name__ == '__main__':
    create_radiation_license_report()
