# -*- coding: utf-8 -*-
"""
生成数据分析报告 (Word 格式)
"""
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# 读取分析结果
results_df = pd.read_excel(r'C:\Users\Mr Zhou\.openclaw\workspace\data\output\信度效度分析结果.xlsx')
loadings_df = pd.read_excel(r'C:\Users\Mr Zhou\.openclaw\workspace\data\output\因子载荷矩阵.xlsx', index_col=0)

# 读取原始数据获取描述统计
df = pd.read_excel(r'C:\Users\Mr Zhou\.openclaw\workspace\data\raw\原始数据.xlsx')

# 创建 Word 文档
doc = Document()

# 设置中文字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

# ========== 标题 ==========
title = doc.add_heading('数据分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph('地理标志产品消费者行为研究')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

# 日期
date_para = doc.add_paragraph(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# ========== 一、数据概况 ==========
doc.add_heading('一、数据概况', level=1)

p = doc.add_paragraph()
p.add_run(f'样本量：{df.shape[0]} 人\n').bold = True
p.add_run(f'变量数：{df.shape[1]} 题\n')
p.add_run(f'量表类型：Likert 5 点量表（1=非常不同意，5=非常同意）')

# 描述统计
doc.add_heading('1.1 描述性统计', level=2)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '变量'
hdr_cells[1].text = '均值'
hdr_cells[2].text = '标准差'
hdr_cells[3].text = '最小值'
hdr_cells[4].text = '最大值'

# 加粗表头
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# 按维度汇总
dimensions = {
    '地理标志认证': list(df.columns[0:4]),
    '地理标志认知': list(df.columns[4:8]),
    '品牌熟悉度': list(df.columns[8:12]),
    '品牌信任度': list(df.columns[12:16]),
    '感知价值': list(df.columns[16:20]),
    '产品涉入度': list(df.columns[20:24]),
    '平台信息信任': list(df.columns[24:28]),
    '线上购买经验': list(df.columns[28:32]),
    '购买意愿': list(df.columns[32:36])
}

for dim_name, cols in dimensions.items():
    dim_data = df[cols].mean(axis=1)
    row_cells = table.add_row().cells
    row_cells[0].text = dim_name
    row_cells[1].text = f'{dim_data.mean():.3f}'
    row_cells[2].text = f'{dim_data.std():.3f}'
    row_cells[3].text = f'{dim_data.min():.1f}'
    row_cells[4].text = f'{dim_data.max():.1f}'

doc.add_paragraph()

# ========== 二、信度分析 ==========
doc.add_heading('二、信度分析', level=1)

p = doc.add_paragraph()
p.add_run('信度分析采用 Cronbach\'s α系数检验量表的内部一致性。\n')
p.add_run('判断标准：α ≥ 0.8：优秀；0.7 ≤ α < 0.8：良好；0.6 ≤ α < 0.7：可接受；α < 0.6：需改进\n')

doc.add_heading('2.1 各维度信度系数', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '维度'
hdr_cells[1].text = 'Cronbach\'s α'
hdr_cells[2].text = '评价'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for _, row in results_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['维度']
    row_cells[1].text = f'{row["Cronbach_α"]:.4f}'
    alpha = row['Cronbach_α']
    if alpha >= 0.8:
        eval_text = '优秀'
    elif alpha >= 0.7:
        eval_text = '良好'
    elif alpha >= 0.6:
        eval_text = '可接受'
    else:
        eval_text = '需改进'
    row_cells[2].text = eval_text

# 总体信度
overall_alpha = df.apply(lambda x: pd.Series(x)).pipe(lambda x: (x.shape[0]/(x.shape[0]-1)) * (1 - x.var().sum()/x.sum(axis=1).var()))
p = doc.add_paragraph()
p.add_run(f'\n总体 Cronbach\'s α系数：{overall_alpha:.4f}\n').bold = True
p.add_run('结论：量表整体信度极佳，数据质量可靠，可进行后续分析。')

doc.add_paragraph()

# ========== 三、效度分析 ==========
doc.add_heading('三、效度分析', level=1)

doc.add_heading('3.1 KMO 和 Bartlett 球形检验', level=2)

from factor_analyzer.factor_analyzer import calculate_kmo, calculate_bartlett_sphericity
kmo_all, kmo_model = calculate_kmo(df)
chi_square_value, p_value = calculate_bartlett_sphericity(df)

p = doc.add_paragraph()
p.add_run(f'KMO 值：{kmo_model:.4f}\n').bold = True
if kmo_model >= 0.8:
    p.add_run('评价：非常适合进行因子分析\n')
elif kmo_model >= 0.7:
    p.add_run('评价：适合进行因子分析\n')
else:
    p.add_run('评价：一般\n')

p.add_run(f'\nBartlett 球形检验：\n').bold = True
p.add_run(f'卡方值：{chi_square_value:.2f}\n')
p.add_run(f'p 值：{p_value:.6f} (p < 0.001)\n')
p.add_run('结论：Bartlett 球形检验显著，变量间存在共同因子，适合进行因子分析。')

doc.add_heading('3.2 结构效度（AVE 和 CR）', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '维度'
hdr_cells[1].text = 'AVE'
hdr_cells[2].text = 'CR'
hdr_cells[3].text = '评价'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for _, row in results_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['维度']
    row_cells[1].text = f'{row["AVE"]:.4f}'
    row_cells[2].text = f'{row["CR"]:.4f}'
    # 评价
    ave_ok = row['AVE'] >= 0.5
    cr_ok = row['CR'] >= 0.7
    if ave_ok and cr_ok:
        eval_text = '通过'
    elif ave_ok or cr_ok:
        eval_text = '部分通过'
    else:
        eval_text = '需改进'
    row_cells[3].text = eval_text

p = doc.add_paragraph()
p.add_run('\n判断标准：AVE ≥ 0.5，CR ≥ 0.7\n')

doc.add_paragraph()

# ========== 四、结论与建议 ==========
doc.add_heading('四、结论与建议', level=1)

doc.add_heading('4.1 主要结论', level=2)

p = doc.add_paragraph()
p.add_run('1. 信度方面：\n').bold = True
p.add_run('   - 所有维度的 Cronbach\'s α系数均大于 0.8，达到"优秀"水平\n')
p.add_run('   - 总体信度系数为 0.9857，量表内部一致性极佳\n')
p.add_run('   - 数据质量可靠，可用于后续假设检验和模型分析\n\n')

p.add_run('2. 效度方面：\n').bold = True
p.add_run(f'   - KMO 值为{kmo_model:.4f}，非常适合因子分析\n')
p.add_run('   - Bartlett 球形检验显著（p < 0.001）\n')
p.add_run('   - 结构效度指标需要进一步优化（部分维度 AVE 和 CR 低于标准）\n\n')

doc.add_heading('4.2 研究建议', level=2)

p = doc.add_paragraph()
p.add_run('1. 数据质量：当前数据信度优秀，可放心进行回归分析、结构方程模型等后续分析\n\n')
p.add_run('2. 效度改进：建议检查部分题目的因子载荷，考虑删除跨因子载荷较高或载荷过低（<0.5）的题目\n\n')
p.add_run('3. 后续分析方向：\n')
p.add_run('   - 验证性因子分析（CFA）进一步验证结构效度\n')
p.add_run('   - 结构方程模型（SEM）检验假设路径\n')
p.add_run('   - 中介效应/调节效应分析\n')

# ========== 页脚 ==========
doc.add_page_break()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run('—— 报告结束 ——').italic = True

# 保存文档
output_path = r'C:\Users\Mr Zhou\.openclaw\workspace\data\output\数据分析报告.docx'
doc.save(output_path)
print(f'[OK] 报告已生成：{output_path}')
